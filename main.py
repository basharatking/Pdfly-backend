"""
PDFly Backend  ·  FastAPI  ·  v3.0
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Deploy on Replit:
  run = "uvicorn main:app --host 0.0.0.0 --port 8080"

Endpoints:
  GET  /health
  POST /merge-pdf
  POST /split-pdf
  POST /compress-pdf
  POST /rotate-pdf
  POST /add-watermark
  POST /pdf-to-word
  POST /pdf-to-excel
  POST /pdf-to-jpg
  POST /jpg-to-pdf
  POST /unlock-pdf
  POST /ocr-check
"""

import io
import os
import zipfile
import tempfile
from pathlib import Path
from typing import List

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware

import pypdf
import pdfplumber
from PIL import Image
import img2pdf
from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
import fitz  # PyMuPDF
# Compatibility fix
if not hasattr(fitz, 'open'):
    fitz.open = fitz.Document
# ─── App Setup ────────────────────────────────────────────────────
app = FastAPI(
    title="PDFly API",
    version="3.0",
    description="PDFly — Free PDF tools backend — merge, split, compress, convert and more.",
    docs_url="/docs",
    redoc_url="/redoc",
)

# CORS — allow all origins so Vercel frontend can call Replit backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=[
        "X-Original-Size",
        "X-Compressed-Size",
        "X-Savings-Pct",
        "Content-Disposition",
    ],
)

# ─── Config ───────────────────────────────────────────────────────
FREE_LIMIT_MB  = int(os.environ.get("FREE_LIMIT_MB",  "25"))
PREM_LIMIT_MB  = int(os.environ.get("PREM_LIMIT_MB", "200"))
FREE_LIMIT_BYTES = FREE_LIMIT_MB  * 1024 * 1024
PREM_LIMIT_BYTES = PREM_LIMIT_MB  * 1024 * 1024

TMPDIR = Path(tempfile.gettempdir()) / "livepdf"
TMPDIR.mkdir(exist_ok=True)

# ─── Helpers ──────────────────────────────────────────────────────
def stream_file(
    data: bytes,
    media_type: str,
    filename: str,
    extra: dict = None,
) -> StreamingResponse:
    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"',
        "Access-Control-Expose-Headers": (
            "X-Original-Size, X-Compressed-Size, X-Savings-Pct"
        ),
    }
    if extra:
        headers.update(extra)
    return StreamingResponse(io.BytesIO(data), media_type=media_type, headers=headers)


async def read_file(upload: UploadFile, premium: bool = False) -> bytes:
    data = await upload.read()
    limit = PREM_LIMIT_BYTES if premium else FREE_LIMIT_BYTES
    if len(data) > limit:
        mb = PREM_LIMIT_MB if premium else FREE_LIMIT_MB
        raise HTTPException(
            413,
            f"File too large. {'Premium' if premium else 'Free'} plan limit is {mb} MB.",
        )
    if not data:
        raise HTTPException(400, "Uploaded file is empty.")
    return data


def stem(filename: str) -> str:
    return Path(filename or "file").stem

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ROUTES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/", tags=["Info"])
def root():
    return {
        "name": "PDFly API",
        "version": "3.0",
        "status": "running",
        "docs": "/docs",
        "tools": [
            "/merge-pdf", "/split-pdf", "/compress-pdf",
            "/rotate-pdf", "/add-watermark", "/pdf-to-word",
            "/pdf-to-excel", "/pdf-to-jpg", "/jpg-to-pdf",
            "/unlock-pdf", "/ocr-check",
        ],
    }


@app.get("/health", tags=["Info"])
def health():
    return {"status": "ok", "version": "3.0"}


# ── 1. MERGE PDF ─────────────────────────────────────────────────
@app.post("/merge-pdf", tags=["Tools"])
async def merge_pdf(
    files: List[UploadFile] = File(..., description="2 or more PDF files"),
    premium: bool = Form(False),
):
    """Combine multiple PDF files into one document."""
    if len(files) < 2:
        raise HTTPException(400, "Please provide at least 2 PDF files to merge.")

    writer = pypdf.PdfWriter()
    for upload in files:
        data = await read_file(upload, premium)
        try:
            reader = pypdf.PdfReader(io.BytesIO(data))
            for page in reader.pages:
                writer.add_page(page)
        except Exception as e:
            raise HTTPException(400, f"Could not read '{upload.filename}': {e}")

    out = io.BytesIO()
    writer.write(out)
    return stream_file(out.getvalue(), "application/pdf", "merged.pdf")


# ── 2. SPLIT PDF ─────────────────────────────────────────────────
@app.post("/split-pdf", tags=["Tools"])
async def split_pdf(
    file: UploadFile = File(...),
    mode: str = Form("each"),          # "each" | "range"
    start_page: int = Form(1),
    end_page: int = Form(1),
    premium: bool = Form(False),
):
    """Split PDF into individual pages (ZIP) or extract a page range."""
    data = await read_file(file, premium)
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        total = len(reader.pages)
    except Exception as e:
        raise HTTPException(400, f"Could not read PDF: {e}")

    if mode == "range":
        s = max(1, start_page) - 1
        e = min(total, end_page)
        if s >= e:
            raise HTTPException(400, f"Invalid page range. PDF has {total} pages.")
        writer = pypdf.PdfWriter()
        for i in range(s, e):
            writer.add_page(reader.pages[i])
        out = io.BytesIO()
        writer.write(out)
        fname = f"{stem(file.filename)}_pages_{s+1}-{e}.pdf"
        return stream_file(out.getvalue(), "application/pdf", fname)

    # mode == "each" → ZIP
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i in range(total):
            w = pypdf.PdfWriter()
            w.add_page(reader.pages[i])
            pb = io.BytesIO()
            w.write(pb)
            zf.writestr(f"page_{str(i+1).zfill(3)}.pdf", pb.getvalue())
    return stream_file(zip_buf.getvalue(), "application/zip", "split_pages.zip")


# ── 3. COMPRESS PDF ──────────────────────────────────────────────
@app.post("/compress-pdf", tags=["Tools"])
async def compress_pdf(
    file: UploadFile = File(...),
    level: str = Form("medium"),    # "low" | "medium" | "high"
    premium: bool = Form(False),
):
    """Compress PDF by recompressing embedded images. Returns before/after sizes in headers."""
    data = await read_file(file, premium)
    orig_size = len(data)

    quality_map = {"low": 85, "medium": 60, "high": 35}
    q = quality_map.get(level, 60)
    max_dim_map = {"low": 1600, "medium": 1200, "high": 900}
    max_dim = max_dim_map.get(level, 1200)

    try:
        doc = fitz.open(stream=data, filetype="pdf")
        for page in doc:
            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    base = doc.extract_image(xref)
                    pil = Image.open(io.BytesIO(base["image"])).convert("RGB")
                    w, h = pil.size
                    if max(w, h) > max_dim:
                        scale = max_dim / max(w, h)
                        pil = pil.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
                    buf = io.BytesIO()
                    pil.save(buf, format="JPEG", quality=q, optimize=True)
                    doc.update_stream(xref, buf.getvalue())
                except Exception:
                    pass  # skip unprocessable images

        out = io.BytesIO()
        doc.save(out, garbage=4, deflate=True, clean=True)
        doc.close()
    except Exception as e:
        raise HTTPException(500, f"Compression failed: {e}")

    comp = out.getvalue()
    comp_size = len(comp)
    savings = round((1 - comp_size / orig_size) * 100, 1) if orig_size else 0

    return stream_file(
        comp, "application/pdf",
        f"compressed_{file.filename}",
        {
            "X-Original-Size":   str(orig_size),
            "X-Compressed-Size": str(comp_size),
            "X-Savings-Pct":     str(savings),
        },
    )


# ── 4. PDF TO WORD ───────────────────────────────────────────────
@app.post("/pdf-to-word", tags=["Tools"])
async def pdf_to_word(
    file: UploadFile = File(...),
    premium: bool = Form(False),
):
    """Extract text from PDF into a .docx Word document."""
    data = await read_file(file, premium)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            total = len(pdf.pages)
            if total == 0:
                raise HTTPException(400, "PDF has no pages.")
            for i, page in enumerate(pdf.pages):
                h = doc.add_heading(f"Page {i + 1}", level=1)
                h.runs[0].font.color.rgb = RGBColor(0xE0, 0x30, 0x20)
                text = page.extract_text() or ""
                if text.strip():
                    for line in text.split("\n"):
                        p = doc.add_paragraph(line)
                        p.paragraph_format.space_after = Pt(2)
                else:
                    p = doc.add_paragraph(
                        "[No extractable text on this page — may be a scanned PDF]"
                    )
                    p.runs[0].italic = True
                if i < total - 1:
                    doc.add_page_break()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Conversion failed: {e}")

    out = io.BytesIO()
    doc.save(out)
    return stream_file(
        out.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{stem(file.filename)}.docx",
    )


# ── 5. PDF TO EXCEL (SMART — single sheet, all pages) ────────────
@app.post("/pdf-to-excel", tags=["Tools"])
async def pdf_to_excel(
    file: UploadFile = File(...),
    mode: str = Form("smart"),   # "smart" | "tables" | "text"
    premium: bool = Form(False),
):
    """
    Extract PDF tables/text to Excel.
    Smart mode: detects tables first, falls back to text. All pages → ONE sheet.
    """
    data = await read_file(file, premium)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.freeze_panes = "A2"

    # Styles
    hdr_fill  = PatternFill("solid", fgColor="E03020")
    hdr_font  = Font(bold=True, color="FFFFFF", size=11)
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ts        = Side(style="thin", color="DDDDDD")
    tbdr      = Border(left=ts, right=ts, top=ts, bottom=ts)
    alt_fill  = PatternFill("solid", fgColor="FFF8F7")

    cur_row      = 1
    hdrs_written = False
    rows_written = 0

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):

                # ── Table extraction ──
                if mode in ("smart", "tables"):
                    tables = page.extract_tables() or []
                    for tbl in tables:
                        if not tbl:
                            continue
                        raw_hdrs = [str(c or "").strip() for c in tbl[0]]

                        # Write header row once
                        if not hdrs_written:
                            for ci, h in enumerate(raw_hdrs, start=1):
                                cell = ws.cell(row=1, column=ci, value=h)
                                cell.fill      = hdr_fill
                                cell.font      = hdr_font
                                cell.alignment = hdr_align
                                cell.border    = tbdr
                            cur_row      = 2
                            hdrs_written = True

                        # Append data rows
                        for row in tbl[1:]:
                            if not any(v for v in row):
                                continue
                            fill = alt_fill if cur_row % 2 == 0 else None
                            for ci, val in enumerate(row, start=1):
                                c = ws.cell(row=cur_row, column=ci,
                                            value=str(val or "").strip())
                                c.border = tbdr
                                if fill:
                                    c.fill = fill
                            cur_row      += 1
                            rows_written += 1

                # ── Text mode (fallback or explicit) ──
                if mode == "text" or (mode == "smart" and not hdrs_written):
                    text = page.extract_text() or ""
                    if not text.strip():
                        continue
                    lines = [l for l in text.split("\n") if l.strip()]
                    if not hdrs_written:
                        for ci, h in enumerate(["Page", "Line", "Text"], start=1):
                            c = ws.cell(row=1, column=ci, value=h)
                            c.fill = hdr_fill; c.font = hdr_font
                            c.alignment = hdr_align; c.border = tbdr
                        cur_row = 2; hdrs_written = True
                    for ln, line in enumerate(lines, start=1):
                        ws.cell(row=cur_row, column=1, value=page_num)
                        ws.cell(row=cur_row, column=2, value=ln)
                        ws.cell(row=cur_row, column=3, value=line)
                        cur_row += 1; rows_written += 1

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Excel conversion failed: {e}")

    # Auto column widths
    for col in ws.columns:
        max_len   = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                max_len = max(max_len, len(str(cell.value or "")))
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(max(max_len + 3, 12), 55)

    if rows_written == 0:
        ws.cell(row=2, column=1,
                value="No tables found. Try switching to 'text' mode.")

    out = io.BytesIO()
    wb.save(out)
    return stream_file(
        out.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        f"{stem(file.filename)}.xlsx",
    )


# ── 6. PDF TO JPG ────────────────────────────────────────────────
@app.post("/pdf-to-jpg", tags=["Tools"])
async def pdf_to_jpg(
    file: UploadFile = File(...),
    dpi: int = Form(150),
    premium: bool = Form(False),
):
    """Render each PDF page as a JPG image. Returns a ZIP archive."""
    data = await read_file(file, premium)
    dpi  = max(72, min(dpi, 300))

    try:
        doc = fitz.open(stream=data, filetype="pdf")
        mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for i, page in enumerate(doc):
                pix = page.get_pixmap(matrix=mat, alpha=False)
                zf.writestr(f"page_{str(i+1).zfill(3)}.jpg", pix.tobytes("jpeg"))
        doc.close()
    except Exception as e:
        raise HTTPException(500, f"Conversion failed: {e}")

    return stream_file(
        zip_buf.getvalue(), "application/zip",
        f"{stem(file.filename)}_images.zip",
    )


# ── 7. JPG / PNG / TXT → PDF ─────────────────────────────────────
@app.post("/jpg-to-pdf", tags=["Tools"])
async def jpg_to_pdf(
    files: List[UploadFile] = File(...),
    premium: bool = Form(False),
):
    """Convert JPG/PNG images (or TXT files) into a PDF document."""
    if not files:
        raise HTTPException(400, "No files provided.")

    images    = []
    txt_pages = []

    for upload in files:
        raw = await read_file(upload, premium)
        ext = Path(upload.filename or "").suffix.lower()
        if ext in {".jpg", ".jpeg", ".png", ".webp"}:
            img = Image.open(io.BytesIO(raw)).convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=92)
            images.append(buf.getvalue())
        elif ext == ".txt":
            txt_pages.append(raw.decode("utf-8", errors="replace"))
        else:
            raise HTTPException(400, f"Unsupported file type: {upload.filename}")

    out = io.BytesIO()
    try:
        if images and not txt_pages:
            out.write(img2pdf.convert(images))
        else:
            c = rl_canvas.Canvas(out, pagesize=A4)
            W, H = A4
            for ib in images:
                pi = Image.open(io.BytesIO(ib))
                iw, ih = pi.size
                r  = min(W / iw, H / ih, 1.0)
                dw, dh = iw * r, ih * r
                c.drawImage(ImageReader(io.BytesIO(ib)),
                            (W - dw) / 2, (H - dh) / 2, dw, dh)
                c.showPage()
            for txt in txt_pages:
                c.setFont("Helvetica", 11)
                m, lh, y = 50, 16, H - 50
                for line in txt.split("\n"):
                    if y < 66:
                        c.showPage()
                        c.setFont("Helvetica", 11)
                        y = H - 50
                    c.drawString(m, y, line[:110])
                    y -= lh
                c.showPage()
            c.save()
    except Exception as e:
        raise HTTPException(500, f"PDF creation failed: {e}")

    return stream_file(out.getvalue(), "application/pdf", "converted.pdf")


# ── 8. ROTATE PDF ────────────────────────────────────────────────
@app.post("/rotate-pdf", tags=["Tools"])
async def rotate_pdf(
    file: UploadFile = File(...),
    angle: int = Form(90),           # 90 | 180 | 270
    pages: str = Form("all"),        # "all" | "odd" | "even"
    premium: bool = Form(False),
):
    """Rotate PDF pages by 90, 180, or 270 degrees."""
    data = await read_file(file, premium)
    if angle not in (90, 180, 270):
        raise HTTPException(400, "Angle must be 90, 180, or 270.")
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        for i, page in enumerate(doc):
            pn = i + 1
            if pages == "odd"  and pn % 2 == 0: continue
            if pages == "even" and pn % 2 != 0: continue
            page.set_rotation((page.rotation + angle) % 360)
        out = io.BytesIO()
        doc.save(out)
        doc.close()
    except Exception as e:
        raise HTTPException(500, f"Rotation failed: {e}")

    return stream_file(out.getvalue(), "application/pdf",
                       f"rotated_{file.filename}")


# ── 9. UNLOCK PDF ────────────────────────────────────────────────
@app.post("/unlock-pdf", tags=["Tools"])
async def unlock_pdf(
    file: UploadFile = File(...),
    password: str = Form(""),
    premium: bool = Form(False),
):
    """Remove password encryption from a PDF file."""
    data = await read_file(file, premium)
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        if doc.is_encrypted:
            if not doc.authenticate(password):
                raise HTTPException(400, "Incorrect password. Could not unlock PDF.")
        out = io.BytesIO()
        doc.save(out, encryption=fitz.PDF_ENCRYPT_NONE)
        doc.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Unlock failed: {e}")

    return stream_file(out.getvalue(), "application/pdf",
                       f"unlocked_{file.filename}")


# ── 10. ADD WATERMARK ────────────────────────────────────────────
@app.post("/add-watermark", tags=["Tools"])
async def add_watermark(
    file: UploadFile = File(...),
    text: str = Form("CONFIDENTIAL"),
    opacity: float = Form(0.2),
    position: str = Form("center"),  # "center" | "top" | "bottom"
    premium: bool = Form(False),
):
    """Add a text watermark to every page of a PDF."""
    if not text.strip():
        raise HTTPException(400, "Watermark text cannot be empty.")
    opacity = max(0.05, min(opacity, 0.95))

    data = await read_file(file, premium)
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        for page in doc:
            w, h  = page.rect.width, page.rect.height
            fs    = min(w, h) * 0.08
            color = (0.55, 0.55, 0.55)
            if position == "center":
                page.insert_text(
                    fitz.Point(w * 0.15, h * 0.55), text,
                    fontsize=fs, rotate=45,
                    color=color, fill_opacity=opacity, overlay=True,
                )
            elif position == "top":
                page.insert_text(
                    fitz.Point(w * 0.5 - len(text) * fs * 0.25, h - fs - 20), text,
                    fontsize=fs, color=color, fill_opacity=opacity, overlay=True,
                )
            else:  # bottom
                page.insert_text(
                    fitz.Point(w * 0.5 - len(text) * fs * 0.25, fs + 20), text,
                    fontsize=fs, color=color, fill_opacity=opacity, overlay=True,
                )
        out = io.BytesIO()
        doc.save(out)
        doc.close()
    except Exception as e:
        raise HTTPException(500, f"Watermark failed: {e}")

    return stream_file(out.getvalue(), "application/pdf",
                       f"watermarked_{file.filename}")


# ── 11. OCR CHECK ────────────────────────────────────────────────
@app.post("/ocr-check", tags=["Tools"])
async def ocr_check(
    file: UploadFile = File(...),
    premium: bool = Form(False),
):
    """
    Detect if a PDF is scanned (image-based) or has a text layer.
    Returns JSON with is_scanned, avg_chars_per_page, total_pages, message.
    """
    data = await read_file(file, premium)
    total_chars = 0
    total_pages = 0
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            total_pages = len(pdf.pages)
            # Sample first 5 pages for speed
            for page in pdf.pages[:5]:
                total_chars += len((page.extract_text() or "").strip())
    except Exception as e:
        raise HTTPException(400, f"Could not read PDF: {e}")

    avg = total_chars / max(total_pages, 1)
    is_scanned = avg < 50  # < 50 chars/page → likely image-based

    return JSONResponse({
        "is_scanned":         is_scanned,
        "avg_chars_per_page": round(avg, 1),
        "total_pages":        total_pages,
        "message": (
            "Scanned PDF detected. Text extraction may not work — OCR required."
            if is_scanned else
            "PDF has a text layer and should convert accurately."
        ),
    })
